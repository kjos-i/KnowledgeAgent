"""Output artifacts — render an AgentAnswer or a chat transcript to files.

Backend + framework-agnostic (no Flet, no GUI state): pure renderers plus a
multi-format `save_answer` / `save_chat`, shared by the GUI Save buttons and
the CLI `--output` flag. The GUI owns only the destination picker; the actual
render + write lives here so both front-ends stay consistent and neither
re-implements it.

Formats:
  - md    Markdown report — headings, question, answer, sources (figures inline)
  - txt   Plain text — the same content with the Markdown markers stripped
  - docx  Word document (python-docx) — headings + answer + sources
  - json  Lossless AgentAnswer dump (answers only; re-loadable into AgentAnswer)

Filenames are `<timestamp>_<query-slug>.<ext>` (answers) and
`<timestamp>_chat_<slug>.<ext>` (chats), so every format written in one save
shares a stem and a results folder stays chronological + human-scannable.

Promoted from the old `gui/artifacts.py` so the CLI can save too; the saved-file
layout for md/json is unchanged.
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

from langchain_core.messages import AIMessage, BaseMessage, HumanMessage

if TYPE_CHECKING:
    from collections.abc import Iterable

    from knowledge_agent.models import AgentAnswer, ChunkSource

# Formats offered per artifact — single source for the GUI checkboxes AND the
# CLI --format choices. The order here is the canonical write order.
ANSWER_FORMATS: tuple[str, ...] = ("md", "txt", "docx", "json")
CHAT_FORMATS: tuple[str, ...] = ("md", "txt", "docx")
FORMAT_LABELS: dict[str, str] = {
    "md": "Markdown (.md)",
    "txt": "Plain text (.txt)",
    "docx": "Word (.docx)",
    "json": "JSON (.json)",
}


class SaveError(Exception):
    """Raised when a save can't write to disk — keeps the save APIs honest."""


# ===========================================================================
# Shared shape helpers — one source-line / role shape, rendered per format.
# ===========================================================================


def _chunk_source_body(c: ChunkSource, *, code: bool) -> str:
    """The descriptive part of a chunk-source line (no `[i]` index prefix).

    `code=True` wraps ids in Markdown backticks (for .md); `code=False` is
    plain text (for .txt / .docx). One source for the source-line shape.
    """

    def wrap(s: str) -> str:
        return f"`{s}`" if code else s

    if c.content_type == "figure":
        if c.page is not None:
            return f"Figure at page {c.page} of {wrap(c.doc_id)} · chunk: {wrap(c.chunk_id)}"
        return f"Image: {wrap(c.doc_id)} · chunk: {wrap(c.chunk_id)}"
    return f"doc: {wrap(c.doc_id)} · chunk: {wrap(c.chunk_id)}"


def _role_of(message: BaseMessage) -> str:
    """Human/assistant/other role label for a chat message."""
    if isinstance(message, HumanMessage):
        return "you"
    if isinstance(message, AIMessage):
        return "assistant"
    return type(message).__name__


def _content_of(message: BaseMessage) -> str:
    return message.content if isinstance(message.content, str) else str(message.content)


_SLUG_NONALNUM = re.compile(r"[^a-z0-9]+")


def _slugify(text: str, max_length: int = 50) -> str:
    """Turn free text into a filesystem-safe slug (lowercased, hyphenated).

    Trailing hyphens are stripped; an empty result falls back to "answer" so
    the caller always gets a non-empty stem.
    """
    slug = _SLUG_NONALNUM.sub("-", text.lower()).strip("-")
    if len(slug) > max_length:
        slug = slug[:max_length].rsplit("-", 1)[0] or slug[:max_length]
    return slug or "answer"


# ===========================================================================
# Answer renderers — md / txt / docx.
# ===========================================================================


_KG_PROVENANCE_KEYS = ("doc_id", "chunk_id", "evidence_span")


def _kg_source_provenance(row: dict | None) -> list[tuple[str, str]]:
    """Extract (label, value) provenance pairs from a cited KG row.

    Surfaces doc_id / chunk_id / evidence_span when the Cypher row carried
    them, tolerating RETURN aliases such as `r.doc_id`. These IDs join the
    graph finding back to its LanceDB source chunk. Empty list when the row
    has no recognizable provenance column (or no row was captured).
    """
    if not row:
        return []
    found: dict[str, str] = {}
    for want in _KG_PROVENANCE_KEYS:
        for key, val in row.items():
            if val is None:
                continue
            if key.split(".")[-1].lower() == want and want not in found:
                found[want] = str(val)
    return [(key, found[key]) for key in _KG_PROVENANCE_KEYS if key in found]


def render_answer_markdown(answer: AgentAnswer, query: str) -> str:
    """Render an AgentAnswer (text + sources) as Markdown for display / disk.

    The two source lists (`chunk_sources` from LanceDB; `kg_sources` from
    Neo4j) render as two sections; the `[1]` / `[K0]` markers in the answer
    text reference them by index.
    """
    lines: list[str] = ["# Knowledge Agent Answer", "", f"**Question:** {query}", "", "---", ""]
    if answer.answer.strip():
        lines.append(answer.answer)
    else:
        # direct_retrieval leaves the answer empty (synthesizer skipped) — say
        # so instead of a blank body; the raw chunks follow in Sources.
        lines.append(
            "_(direct retrieve — synthesizer skipped; raw retrieved chunks are listed below)_"
        )
    lines += [
        "",
        "---",
        "",
        f"## Sources ({len(answer.chunk_sources)} chunk, {len(answer.kg_sources)} KG)",
        "",
    ]
    if not answer.chunk_sources and not answer.kg_sources:
        lines.append("_(none)_")
        return "\n".join(lines)
    if answer.chunk_sources:
        lines += ["### Chunk sources", ""]
        for i, c in enumerate(answer.chunk_sources, start=1):
            lines.append(f"**[{i}]** {_chunk_source_body(c, code=True)}")
            if c.content_type == "figure" and c.image_ref:
                # Embed the figure so the saved .md opens with the picture
                # visible (offline reading of cited figures).
                lines += ["", f"![figure {i}]({c.image_ref})"]
            if c.quote:
                label = "caption / OCR" if c.content_type == "figure" else "quote"
                lines += ["", f"> _{label}:_ {c.quote}"]
            lines.append("")
    if answer.kg_sources:
        lines += ["### KG sources", ""]
        for k in answer.kg_sources:
            lines.append(f"**[K{k.hit_index}]**")
            for label, val in _kg_source_provenance(k.row):
                lines += ["", f"> _{label}:_ {val}"]
            if k.quote:
                lines += ["", f"> {k.quote}"]
            lines.append("")
    return "\n".join(lines)


def render_answer_txt(answer: AgentAnswer, query: str) -> str:
    """Render an AgentAnswer as plain text — same content as the Markdown,
    with the `#` / `**` / `>` / image markers stripped (for readers who want
    clean text)."""
    lines: list[str] = ["Knowledge Agent Answer", "", f"Question: {query}", ""]
    if answer.answer.strip():
        lines.append(answer.answer)
    else:
        lines.append("(direct retrieve — synthesizer skipped; raw retrieved chunks below)")
    lines += ["", f"Sources ({len(answer.chunk_sources)} chunk, {len(answer.kg_sources)} KG)", ""]
    if not answer.chunk_sources and not answer.kg_sources:
        lines.append("(none)")
        return "\n".join(lines) + "\n"
    if answer.chunk_sources:
        lines.append("Chunk sources:")
        for i, c in enumerate(answer.chunk_sources, start=1):
            lines.append(f"  [{i}] {_chunk_source_body(c, code=False)}")
            if c.content_type == "figure" and c.image_ref:
                lines.append(f"      figure: {c.image_ref}")
            if c.quote:
                label = "caption / OCR" if c.content_type == "figure" else "quote"
                lines.append(f"      {label}: {c.quote}")
        lines.append("")
    if answer.kg_sources:
        lines.append("KG sources:")
        for k in answer.kg_sources:
            suffix = f" {k.quote}" if k.quote else ""
            lines.append(f"  [K{k.hit_index}]{suffix}")
            for label, val in _kg_source_provenance(k.row):
                lines.append(f"      {label}: {val}")
        lines.append("")
    return "\n".join(lines) + "\n"


def render_answer_docx(answer: AgentAnswer, query: str):
    """Render an AgentAnswer as a python-docx Document (the caller saves it).

    Returns a `docx.document.Document`. python-docx is imported lazily so this
    module stays importable (md / txt / json work) even if docx isn't present.
    """
    from docx import Document

    doc = Document()
    doc.add_heading("Knowledge Agent Answer", level=0)
    p = doc.add_paragraph()
    p.add_run("Question: ").bold = True
    p.add_run(query)
    if answer.answer.strip():
        doc.add_paragraph(answer.answer)
    else:
        doc.add_paragraph().add_run(
            "(direct retrieve — synthesizer skipped; raw retrieved chunks below)"
        ).italic = True
    doc.add_heading(
        f"Sources ({len(answer.chunk_sources)} chunk, {len(answer.kg_sources)} KG)", level=1
    )
    if not answer.chunk_sources and not answer.kg_sources:
        doc.add_paragraph("(none)")
        return doc
    if answer.chunk_sources:
        doc.add_heading("Chunk sources", level=2)
        for i, c in enumerate(answer.chunk_sources, start=1):
            p = doc.add_paragraph()
            p.add_run(f"[{i}] ").bold = True
            p.add_run(_chunk_source_body(c, code=False))
            if c.content_type == "figure" and c.image_ref:
                _embed_figure_docx(doc, c.image_ref, i)
            if c.quote:
                label = "caption / OCR" if c.content_type == "figure" else "quote"
                doc.add_paragraph().add_run(f"{label}: {c.quote}").italic = True
    if answer.kg_sources:
        doc.add_heading("KG sources", level=2)
        for k in answer.kg_sources:
            doc.add_paragraph().add_run(f"[K{k.hit_index}]").bold = True
            for label, val in _kg_source_provenance(k.row):
                doc.add_paragraph().add_run(f"{label}: {val}").italic = True
            if k.quote:
                doc.add_paragraph().add_run(k.quote).italic = True
    return doc


def _embed_figure_docx(doc, image_ref: str, index: int) -> None:
    """Embed a cited figure image into the docx, or note it if unreadable.

    Best-effort: a missing/corrupt image must never fail the whole save.
    """
    from docx.shared import Inches

    img = Path(image_ref)
    if not img.is_file():
        doc.add_paragraph(f"[figure {index} — image not found: {image_ref}]")
        return
    try:
        doc.add_picture(str(img), width=Inches(5.5))
    except Exception:  # any image error degrades to a text note, never a failed save
        doc.add_paragraph(f"[figure {index} — could not embed: {image_ref}]")


# ===========================================================================
# Chat renderers — md / txt / docx.
# ===========================================================================


def render_chat_markdown(messages: list[BaseMessage], timestamp: str) -> str:
    """Render a conversation as a Markdown transcript."""
    lines: list[str] = [f"# Knowledge Agent Chat — {timestamp}", ""]
    for m in messages:
        lines += [f"**{_role_of(m)}**", "", _content_of(m), ""]
    return "\n".join(lines)


def render_chat_txt(messages: list[BaseMessage], timestamp: str) -> str:
    """Render a conversation as a plain-text transcript (no Markdown markers)."""
    lines: list[str] = [f"Knowledge Agent Chat — {timestamp}", ""]
    for m in messages:
        lines += [f"{_role_of(m)}:", _content_of(m), ""]
    return "\n".join(lines) + "\n"


def render_chat_docx(messages: list[BaseMessage], timestamp: str):
    """Render a conversation as a python-docx Document (the caller saves it)."""
    from docx import Document

    doc = Document()
    doc.add_heading(f"Knowledge Agent Chat — {timestamp}", level=0)
    for m in messages:
        doc.add_paragraph().add_run(_role_of(m)).bold = True
        doc.add_paragraph(_content_of(m))
    return doc


# ===========================================================================
# Multi-format save.
# ===========================================================================


def _normalize_formats(formats: Iterable[str], allowed: tuple[str, ...]) -> list[str]:
    """Dedupe + validate `formats` against `allowed`, returned in `allowed`'s
    canonical order. Raises SaveError if none are valid."""
    requested = set(formats)
    chosen = [f for f in allowed if f in requested]
    if not chosen:
        raise SaveError(f"no valid formats in {sorted(requested)}; allowed: {', '.join(allowed)}")
    return chosen


def save_answer(
    answer: AgentAnswer, query: str, results_dir: Path, formats: Iterable[str]
) -> list[Path]:
    """Save `answer` to `results_dir` in each of `formats` (see `ANSWER_FORMATS`).

    All formats share one stem `<timestamp>_<query-slug>` so a save produces a
    named set (`.md`, `.docx`, …). Returns the written paths in canonical
    order. Raises `SaveError` on any filesystem failure.
    """
    chosen = _normalize_formats(formats, ANSWER_FORMATS)
    stem = f"{datetime.now().strftime('%Y-%m-%d_%H%M%S')}_{_slugify(query)}"
    written: list[Path] = []
    try:
        results_dir.mkdir(parents=True, exist_ok=True)
        for fmt in chosen:
            path = results_dir / f"{stem}.{fmt}"
            _write_answer(fmt, answer, query, path)
            written.append(path)
    except OSError as exc:
        raise SaveError(f"could not write to {results_dir}: {exc}") from exc
    return written


def _write_answer(fmt: str, answer: AgentAnswer, query: str, path: Path) -> None:
    if fmt == "md":
        path.write_text(render_answer_markdown(answer, query), encoding="utf-8")
    elif fmt == "txt":
        path.write_text(render_answer_txt(answer, query), encoding="utf-8")
    elif fmt == "json":
        path.write_text(answer.model_dump_json(indent=2), encoding="utf-8")
    elif fmt == "docx":
        render_answer_docx(answer, query).save(str(path))


def save_chat(
    messages: list[BaseMessage], query: str | None, results_dir: Path, formats: Iterable[str]
) -> list[Path]:
    """Save the conversation to `results_dir` in each of `formats` (see
    `CHAT_FORMATS`). `query` is used for the filename slug only. Returns the
    written paths; raises `SaveError` on any filesystem failure."""
    chosen = _normalize_formats(formats, CHAT_FORMATS)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    slug = _slugify(query) if query else "chat"
    stem = f"{timestamp}_chat_{slug}"
    written: list[Path] = []
    try:
        results_dir.mkdir(parents=True, exist_ok=True)
        for fmt in chosen:
            path = results_dir / f"{stem}.{fmt}"
            _write_chat(fmt, messages, timestamp, path)
            written.append(path)
    except OSError as exc:
        raise SaveError(f"could not write to {results_dir}: {exc}") from exc
    return written


def _write_chat(fmt: str, messages: list[BaseMessage], timestamp: str, path: Path) -> None:
    if fmt == "md":
        path.write_text(render_chat_markdown(messages, timestamp), encoding="utf-8")
    elif fmt == "txt":
        path.write_text(render_chat_txt(messages, timestamp), encoding="utf-8")
    elif fmt == "docx":
        render_chat_docx(messages, timestamp).save(str(path))
