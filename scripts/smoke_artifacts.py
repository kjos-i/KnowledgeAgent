"""Smoke test for knowledge_agent.artifacts - render + save an answer & chat
in every format, then reopen the .docx and round-trip the .json to prove they
are real, valid files.

Builds a realistic AgentAnswer (answer text + a text chunk source + a FIGURE
chunk with a real embedded image + a KG source) and a short chat, saves each
in all formats (answer: md/txt/docx/json; chat: md/txt/docx) to a fresh temp
folder, then self-checks: the .docx reopens with the figure embedded, the
.json round-trips back into an AgentAnswer, and the .txt is marker-free.

No DB, no LLM, no keys - pure rendering. Run from the project root:
    python scripts/smoke_artifacts.py

Automated counterpart:
  tests/unit/test_artifacts.py  (all renderers, multi-format save, docx figure
                                 embed, json round-trip, error paths)
Run via `pytest tests/unit/test_artifacts.py`.
"""

from __future__ import annotations

import json
import struct
import tempfile
import zlib
from pathlib import Path

from docx import Document
from langchain_core.messages import AIMessage, HumanMessage

from knowledge_agent.artifacts import (
    ANSWER_FORMATS,
    CHAT_FORMATS,
    save_answer,
    save_chat,
)
from knowledge_agent.models import AgentAnswer, ChunkSource, KGSource


def _write_tiny_png(path: Path) -> None:
    """Write a minimal but VALID 1x1 red PNG (docx-embeddable)."""

    def chunk(typ: bytes, data: bytes) -> bytes:
        body = typ + data
        return (
            struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)
        )

    path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00"))
        + chunk(b"IEND", b"")
    )


def main() -> None:
    out = Path(tempfile.mkdtemp(prefix="ka_smoke_artifacts_"))
    fig = out / "figure.png"
    _write_tiny_png(fig)

    answer = AgentAnswer(
        answer="ESCRT-III drives membrane scission [1]; the spiral is shown in the figure [2]; VPS4 regulates it [K0].",
        chunk_sources=[
            ChunkSource(chunk_id="doc1#3", doc_id="doc1", quote="membrane scission by ESCRT-III"),
            ChunkSource(
                chunk_id="doc1#7",
                doc_id="doc1",
                content_type="figure",
                image_ref=str(fig),
                page=4,
                quote="Figure 2: the ESCRT-III spiral",
            ),
        ],
        kg_sources=[KGSource(hit_index=0, quote="VPS4 -[:REGULATES]-> ESCRT-III")],
    )
    chat = [
        HumanMessage(content="How does ESCRT-III work?"),
        AIMessage(content="It polymerises into spirals that constrict the membrane neck."),
    ]

    print(f"Output folder: {out}\n")

    answer_paths = save_answer(answer, "How does ESCRT-III drive scission?", out, ANSWER_FORMATS)
    print(f"Answer saved ({len(answer_paths)} formats):")
    for p in answer_paths:
        print(f"  {p.name}  ({p.stat().st_size} bytes)")

    chat_paths = save_chat(chat, "How does ESCRT-III work?", out, CHAT_FORMATS)
    print(f"\nChat saved ({len(chat_paths)} formats):")
    for p in chat_paths:
        print(f"  {p.name}  ({p.stat().st_size} bytes)")

    # ---- self-checks -------------------------------------------------------
    print("\nChecks:")
    docx_path = next(p for p in answer_paths if p.suffix == ".docx")
    doc = Document(str(docx_path))
    ok_docx = len(doc.paragraphs) > 0 and len(doc.inline_shapes) >= 1
    print(
        f"  [{'OK' if ok_docx else 'FAIL'}] answer.docx reopens: "
        f"{len(doc.paragraphs)} paragraphs, {len(doc.inline_shapes)} embedded image(s)"
    )

    json_path = next(p for p in answer_paths if p.suffix == ".json")
    restored = AgentAnswer.model_validate(json.loads(json_path.read_text(encoding="utf-8")))
    ok_json = restored.answer == answer.answer and len(restored.chunk_sources) == 2
    print(f"  [{'OK' if ok_json else 'FAIL'}] answer.json round-trips back into an AgentAnswer")

    txt = next(p for p in answer_paths if p.suffix == ".txt").read_text(encoding="utf-8")
    ok_txt = "**" not in txt and "# " not in txt and "figure:" in txt
    print(
        f"  [{'OK' if ok_txt else 'FAIL'}] answer.txt is plain (no markdown markers) + lists the figure"
    )

    print(f"\nOpen the folder to eyeball the files:\n  {out}")
    if not (ok_docx and ok_json and ok_txt):
        raise SystemExit("\nSMOKE FAILED - see checks above")
    print("\nAll checks passed.")


if __name__ == "__main__":
    main()
