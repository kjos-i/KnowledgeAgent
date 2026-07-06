"""Tests for `knowledge_agent.artifacts` — md/txt/docx/json rendering + save.

Pure-function + file-IO tests (tmp_path); no Flet involvement. docx tests
reopen the written file with python-docx and read its text back.
"""

from __future__ import annotations

import json
import struct
import zlib
from pathlib import Path

import pytest
from langchain_core.messages import AIMessage, HumanMessage, SystemMessage

from knowledge_agent.artifacts import (
    ANSWER_FORMATS,
    CHAT_FORMATS,
    SaveError,
    _slugify,
    render_answer_markdown,
    render_answer_txt,
    render_chat_markdown,
    render_chat_txt,
    save_answer,
    save_chat,
)
from knowledge_agent.models import AgentAnswer, ChunkSource, KGSource


def _answer(text: str, *, chunks: int = 0, kgs: int = 0) -> AgentAnswer:
    return AgentAnswer(
        answer=text,
        chunk_sources=[ChunkSource(chunk_id=f"d{i}#{i}", doc_id=f"d{i}") for i in range(chunks)],
        kg_sources=[KGSource(hit_index=i) for i in range(kgs)],
    )


def _docx_text(path: Path) -> str:
    """Concatenated paragraph text of a saved .docx (headings included)."""
    from docx import Document

    return "\n".join(p.text for p in Document(str(path)).paragraphs)


def _tiny_png() -> bytes:
    """A minimal but VALID 1x1 PNG (correct CRCs + real IDAT) for the docx
    figure-embed SUCCESS path — an over-minimal base64 PNG trips python-docx's
    stricter parser (UnexpectedEndOfFileError), so we build a real one."""

    def chunk(typ: bytes, data: bytes) -> bytes:
        body = typ + data
        return (
            struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)  # 1x1, 8-bit RGB
    raw = b"\x00\xff\x00\x00"  # one scanline: filter byte 0 + one red pixel
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


_PNG_1x1 = _tiny_png()


# ---- format constants ----


def test_format_constants():
    assert ANSWER_FORMATS == ("md", "txt", "docx", "json")
    assert CHAT_FORMATS == ("md", "txt", "docx")  # chat has no json


# ---- _slugify ----


def test_slugify_lowercases_and_hyphenates():
    assert _slugify("Hello World") == "hello-world"


def test_slugify_strips_trailing_punctuation():
    assert _slugify("foo bar!?") == "foo-bar"


def test_slugify_empty_falls_back_to_answer():
    assert _slugify("") == "answer"
    assert _slugify("!!!") == "answer"


def test_slugify_respects_max_length():
    assert len(_slugify("word " * 30, max_length=20)) <= 20


# ---- render_answer_markdown ----


def test_render_answer_includes_question_and_body():
    md = render_answer_markdown(_answer("the answer"), "the question")
    assert "the question" in md
    assert "the answer" in md
    assert md.startswith("# Knowledge Agent Answer")


def test_render_answer_empty_body_shows_direct_retrieve_note():
    md = render_answer_markdown(_answer("", chunks=1), "q")
    assert "direct retrieve" in md.lower()
    assert "synthesizer skipped" in md.lower()


def test_render_answer_lists_chunk_and_kg_sources_with_indices():
    md = render_answer_markdown(_answer("body", chunks=2, kgs=2), "q")
    assert "**[1]**" in md and "**[2]**" in md and "Chunk sources" in md
    assert "**[K0]**" in md and "**[K1]**" in md and "KG sources" in md


def test_render_answer_with_no_sources_shows_none():
    assert "_(none)_" in render_answer_markdown(_answer("body"), "q")


# ---- render_answer_txt (plain, no Markdown markers) ----


def test_render_answer_txt_has_question_and_body_no_markers():
    txt = render_answer_txt(_answer("the answer", chunks=1, kgs=1), "the question")
    assert "Question: the question" in txt
    assert "the answer" in txt
    # No Markdown syntax leaks into plain text.
    assert "# " not in txt
    assert "**" not in txt
    assert "![" not in txt
    # Sources still present, just plainly formatted.
    assert "Chunk sources:" in txt
    assert "[1]" in txt
    assert "[K0]" in txt


def test_render_answer_txt_empty_body_note():
    txt = render_answer_txt(_answer("", chunks=1), "q")
    assert "synthesizer skipped" in txt.lower()


# ---- render_answer_docx ----


def test_save_answer_docx_contains_question_and_answer(tmp_path: Path):
    (path,) = save_answer(_answer("the docx answer", chunks=1, kgs=1), "the q", tmp_path, ("docx",))
    assert path.suffix == ".docx"
    text = _docx_text(path)
    assert "Knowledge Agent Answer" in text
    assert "the q" in text
    assert "the docx answer" in text


def test_save_answer_docx_notes_missing_figure_image(tmp_path: Path):
    """A figure chunk whose image file is absent degrades to a text note,
    never a failed save."""
    ans = AgentAnswer(
        answer="body",
        chunk_sources=[
            ChunkSource(chunk_id="c0", doc_id="d0", content_type="figure", image_ref="/no/img.png")
        ],
    )
    (path,) = save_answer(ans, "q", tmp_path, ("docx",))
    assert "image not found" in _docx_text(path)


def test_save_answer_docx_embeds_present_figure(tmp_path: Path):
    """A figure chunk with a real image file → embedded as an inline picture
    (the add_picture SUCCESS path, distinct from the missing-image note)."""
    from docx import Document

    img = tmp_path / "fig.png"
    img.write_bytes(_PNG_1x1)
    ans = AgentAnswer(
        answer="body",
        chunk_sources=[
            ChunkSource(chunk_id="c0", doc_id="d0", content_type="figure", image_ref=str(img))
        ],
    )
    (path,) = save_answer(ans, "q", tmp_path, ("docx",))
    assert len(Document(str(path)).inline_shapes) >= 1  # picture actually embedded


def test_save_chat_docx_contains_roles_and_text(tmp_path: Path):
    """Chat .docx reopens with the role labels + message text (content, not
    just a file-exists check)."""
    (path,) = save_chat(
        [HumanMessage(content="hi there"), AIMessage(content="hello back")],
        "q",
        tmp_path,
        ("docx",),
    )
    text = _docx_text(path)
    assert "you" in text and "assistant" in text
    assert "hi there" in text and "hello back" in text


def test_render_answer_txt_includes_figure_and_quote_lines():
    """Plain-text renders the figure path line + labels a figure quote as
    caption/OCR and a text-chunk quote as quote."""
    ans = AgentAnswer(
        answer="body",
        chunk_sources=[
            ChunkSource(
                chunk_id="c0",
                doc_id="d0",
                content_type="figure",
                image_ref="/p/img.png",
                quote="a caption",
            ),
            ChunkSource(chunk_id="c1", doc_id="d1", quote="a text quote"),
        ],
    )
    txt = render_answer_txt(ans, "q")
    assert "figure: /p/img.png" in txt
    assert "caption / OCR: a caption" in txt
    assert "quote: a text quote" in txt


# ---- save_answer (multi-format) ----


def test_save_answer_writes_all_selected_formats_sharing_a_stem(tmp_path: Path):
    paths = save_answer(_answer("body", chunks=1), "what is X", tmp_path, ANSWER_FORMATS)
    suffixes = {p.suffix for p in paths}
    assert suffixes == {".md", ".txt", ".docx", ".json"}
    # One shared stem across all formats.
    assert len({p.stem for p in paths}) == 1
    assert "what-is-x" in paths[0].stem
    assert all(p.exists() for p in paths)


def test_save_answer_canonical_order_regardless_of_input_order(tmp_path: Path):
    paths = save_answer(_answer("body"), "q", tmp_path, ("json", "md"))
    # Returned in ANSWER_FORMATS order (md before json), not input order.
    assert [p.suffix for p in paths] == [".md", ".json"]


def test_save_answer_json_round_trips_to_agent_answer(tmp_path: Path):
    original = _answer("body", chunks=2, kgs=1)
    (json_path,) = save_answer(original, "q", tmp_path, ("json",))
    restored = AgentAnswer.model_validate(json.loads(json_path.read_text(encoding="utf-8")))
    assert restored.answer == original.answer
    assert len(restored.chunk_sources) == 2
    assert len(restored.kg_sources) == 1


def test_save_answer_rejects_when_no_valid_formats(tmp_path: Path):
    with pytest.raises(SaveError, match="no valid formats"):
        save_answer(_answer("body"), "q", tmp_path, ("pdf", "rtf"))


def test_save_answer_raises_save_error_when_write_fails(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    def boom(*_a, **_k):
        raise OSError("disk full")

    monkeypatch.setattr(Path, "write_text", boom)
    with pytest.raises(SaveError, match="disk full"):
        save_answer(_answer("body"), "q", tmp_path, ("md",))


# ---- chat: render + save ----


def test_render_chat_markdown_renders_roles():
    md = render_chat_markdown(
        [HumanMessage(content="hi"), AIMessage(content="hello"), SystemMessage(content="sys")],
        "2026-07-06_120000",
    )
    assert "**you**" in md and "**assistant**" in md and "SystemMessage" in md


def test_render_chat_txt_plain_roles():
    txt = render_chat_txt([HumanMessage(content="hi"), AIMessage(content="hello")], "ts")
    assert "you:" in txt and "assistant:" in txt
    assert "**" not in txt


def test_save_chat_writes_selected_formats(tmp_path: Path):
    paths = save_chat([HumanMessage(content="hi")], "my q", tmp_path, ("md", "txt", "docx"))
    assert {p.suffix for p in paths} == {".md", ".txt", ".docx"}
    assert all("_chat_" in p.name for p in paths)


def test_save_chat_slug_falls_back_when_query_none(tmp_path: Path):
    (path,) = save_chat([HumanMessage(content="hi")], None, tmp_path, ("md",))
    assert "_chat_" in path.name
